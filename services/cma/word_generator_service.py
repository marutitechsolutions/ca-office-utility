"""
Word Generator Service for CMA / DPR Builder.
Generates editable .docx reports using python-docx.
"""

import os
import io
import time
import uuid
from datetime import datetime
import matplotlib
matplotlib.use('Agg')
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import matplotlib.pyplot as plt
from services.cma.models import CmaProject
from services.cma.narrative_service import NarrativeService
from services.cma.projection_engine_service import ProjectionEngineService

class WordGeneratorService:
    """Service to build professional Word reports for CMA / DPR."""

    @classmethod
    def _set_cell_background(cls, cell, fill):
        """Sets the background (shading) of a table cell using OXML."""
        # fill is hex color string e.g. '0D47A1'
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    @classmethod
    def _prevent_table_split(cls, table):
        """Prevents table rows from breaking across pages and keeps headers with content."""
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)
        
        # Aggressive Keep: Keep every row with the next one, except the last
        for i in range(len(table.rows) - 1):
            for cell in table.rows[i].cells:
                for p in cell.paragraphs:
                    pPr = p._element.get_or_add_pPr()
                    keepNext = OxmlElement('w:keepNext')
                    pPr.append(keepNext)

    @classmethod
    def _set_keep_with_next(cls, p):
        """Forces a paragraph to stay with the following paragraph/table."""
        pPr = p._element.get_or_add_pPr()
        keepNext = OxmlElement('w:keepNext')
        pPr.append(keepNext)

    @classmethod
    def _setup_styles(cls, doc, theme=None):
        """Sets the default font and heading styles based on theme."""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        
        # Heading 1 for Section Titles
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Times New Roman'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.paragraph_format.keep_with_next = True
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(6)
        if theme:
            h1.font.color.rgb = RGBColor(*theme.PRIMARY_RGB)
        else:
            h1.font.color.rgb = RGBColor(13, 71, 161) # Deep Blue fallback

        # Heading 2
        h2 = doc.styles['Heading 2']
        h2.font.name = 'Times New Roman'
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.paragraph_format.keep_with_next = True
        h2.paragraph_format.space_before = Pt(8)
        h1.paragraph_format.space_after = Pt(4)

    @classmethod
    def _fmt_val(cls, val, precision=2) -> str:
        """Safely formats a value for Word tables, avoiding dict/non-numeric crashes."""
        if val is None: return "0.00"
        if isinstance(val, dict): 
            # Handle dictionary wrappers if they contain values
            if "value" in val: val = val["value"]
            elif "total" in val: val = val["total"]
            elif "amount" in val: val = val["amount"]
            else: return "N/A"
        try:
            return f"{float(val):.{precision}f}"
        except (ValueError, TypeError):
            return str(val)

    @classmethod
    def _get_year_headers(cls, projections, project):
        """Build year header labels with audit status tags (Parity with PDF)."""
        headers = []
        for r in projections:
            orig_label = r.get("year_label", "N/A")
            # Reuse the status logic from PDF
            status = ""
            if r.get("is_actual"):
                ds = r.get("data_status")
                if ds:
                    if ds == "Audited": status = " (Aud.)"
                    elif ds == "Provisional": status = " (Prov.)"
                else:
                    for ad in project.audited_history:
                        if ad.year_label == orig_label:
                            if ad.data_type == "Audited": status = " (Aud.)"
                            elif ad.data_type == "Provisional": status = " (Prov.)"
                            break
            else:
                status = " (P)"
            
            headers.append(f"{orig_label}{status}")
        return headers

    @classmethod
    def generate_docx(cls, project: CmaProject, output_path: str) -> str:
        """
        Generates the full premium consultant-grade Word report.
        """
        from services.cma.models import ReportMode
        from services.cma.report_theme import get_theme
        
        report_mode = project.profile.report_mode
        theme = get_theme(report_mode)
        doc = Document()
        cls._setup_styles(doc, theme)
        
        # Financial Data & Validation
        proj_results = ProjectionEngineService.generate_full_projections(project)
        
        # 1. Slice Projections based on Mode
        max_years = 3 if report_mode == ReportMode.LITE.value else 5
        proj_results = proj_results[:max_years]

        # 2. Master Section Flow (Parity with PDF PRO mode)
        cls._add_section_A_cover(doc, project, theme)
        doc.add_page_break()
        
        cls._add_section_B_contents(doc, project)
        doc.add_page_break()
        
        cls._add_section_B1_analytical_profile(doc, project)
        cls._add_section_C_summary(doc, project)
        
        if report_mode == ReportMode.LITE.value:
            # Mode 1: LITE (Compact)
            cls._add_section_DASH_dashboard(doc, project, proj_results, theme)
            cls._add_section_D_snapshot(doc, project, theme)
            cls._add_section_F_promoter(doc, project)
            cls._add_section_H_cost(doc, project, theme)
            cls._add_section_I_finance(doc, project, theme)
            cls._add_section_L_operating_stmt(doc, project, proj_results, theme)
            cls._add_section_M_balance_sheet(doc, project, proj_results, theme)
            cls._add_section_Q_dscr(doc, project, proj_results, theme)
            cls._add_section_AD_readiness(doc, project, proj_results, theme)
            cls._add_section_X_assumptions(doc, project)
        elif report_mode == ReportMode.CMA.value:
            # Mode 3: CMA (Banker Analytical)
            cls._add_section_DASH_dashboard(doc, project, proj_results, theme)
            cls._add_section_E_entity(doc, project, theme)
            cls._add_section_L_operating_stmt(doc, project, proj_results, theme)
            cls._add_section_M_balance_sheet(doc, project, proj_results, theme)
            cls._add_section_N_cash_flow(doc, project, proj_results, theme)
            cls._add_section_W_cma_data(doc, project, proj_results, theme)
            cls._add_section_AC_mpbf(doc, project, proj_results, theme)
            cls._add_section_R_liquidity(doc, project, proj_results, theme)
            cls._add_section_AD_readiness(doc, project, proj_results, theme)
            cls._add_section_X_assumptions(doc, project)
        else:
            # Mode 2: PRO (Full Premium Parity)
            cls._add_section_DASH_dashboard(doc, project, proj_results, theme)
            cls._add_section_D_snapshot(doc, project, theme)
            cls._add_section_E_entity(doc, project, theme)
            cls._add_section_F_promoter(doc, project)
            cls._add_section_G_employment(doc, project, theme)
            cls._add_section_H_cost(doc, project, theme)
            cls._add_section_I_finance(doc, project, theme)
            cls._add_section_J_financial_data(doc, project, proj_results, theme)
            cls._add_section_K_graphics(doc, project, proj_results, theme)
            cls._add_section_L_operating_stmt(doc, project, proj_results, theme)
            cls._add_section_N_cash_flow(doc, project, proj_results, theme)
            cls._add_section_M_balance_sheet(doc, project, proj_results, theme)
            cls._add_section_O_fixed_assets(doc, project, proj_results, theme)
            cls._add_section_V_repayment(doc, project)
            cls._add_section_Q_dscr(doc, project, proj_results, theme)
            cls._add_section_AA_monthly_repayment(doc, project, theme)
            cls._add_section_P_expenses(doc, project, proj_results, theme)
            cls._add_section_R_liquidity(doc, project, proj_results, theme)
            cls._add_section_S_sensitivity(doc, project, proj_results, theme)
            cls._add_section_T_bep(doc, project, proj_results, theme)
            cls._add_section_U_margin(doc, project, proj_results)
            cls._add_section_W_cma_data(doc, project, proj_results, theme)
            cls._add_section_AC_mpbf(doc, project, proj_results, theme)
            cls._add_section_AD_readiness(doc, project, proj_results, theme)
            cls._add_section_X_assumptions(doc, project)
            cls._add_section_Y_security(doc, project)
        
        cls._add_section_Z_declaration(doc, project)
        doc.save(output_path)
        return output_path

    @classmethod
    def _add_section_B1_analytical_profile(cls, doc, project):
        """B1. Analytical Business Profile."""
        h = doc.add_heading("SECTION B1: ANALYTICAL BUSINESS PROFILE", level=1)
        cls._set_keep_with_next(h)
        text = cls._clean_text(NarrativeService.generate_section("business_overview", project))
        doc.add_paragraph(text)
        
        doc.add_heading("Market & Industry Potential", level=2)
        mkt = cls._clean_text(NarrativeService.generate_section("market_potential", project))
        doc.add_paragraph(mkt)

    @classmethod
    def _generate_charts(cls, project, projections):
        """Replicate PDF chart generation for Word."""
        charts = []
        if not projections: return charts
        years = [p['year_label'] for p in projections]
        
        # Chart 1: Revenue
        try:
            plt.clf()
            plt.figure(figsize=(6, 3))
            plt.bar(years, [p['revenue'] for p in projections], color='#0D47A1')
            plt.title('Revenue Growth (Lakhs)', fontsize=10, fontweight='bold')
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=120)
            buf.seek(0)
            charts.append(buf)
            plt.close('all')
        except Exception as e:
            logger.error(f"Error generating revenue chart: {e}")
        
        # Chart 2: PAT
        try:
            plt.clf()
            plt.figure(figsize=(6, 3))
            plt.plot(years, [p['pat'] for p in projections], marker='o', color='#2E7D32', linewidth=2)
            plt.title('Profitability Trend (PAT)', fontsize=10, fontweight='bold')
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=120)
            buf.seek(0)
            charts.append(buf)
            plt.close('all')
        except Exception as e:
            logger.error(f"Error generating PAT chart: {e}")
            
        return charts

    @classmethod
    def _add_picture_safely(cls, doc, buf, width):
        """Helper to add picture with retry logic for WinError 32 (file lock)."""
        max_retries = 3
        for i in range(max_retries):
            try:
                # We use a seek(0) to ensure we're at the start of the buffer
                buf.seek(0)
                doc.add_picture(buf, width=width)
                return True
            except Exception as e:
                if "WinError 32" in str(e) or "used by another process" in str(e):
                    if i < max_retries - 1:
                        time.sleep(0.5) # Wait for lock to release
                        continue
                logger.error(f"Failed to add picture to Word doc: {e}")
                return False
        return False

    @classmethod
    def _add_section_K_graphics(cls, doc, project, projections, theme=None):
        """K. Statistical Graphics."""
        h = doc.add_heading("SECTION K: GRAPHICAL PERFORMANCE ANALYSIS", level=1)
        cls._set_keep_with_next(h)
        charts = cls._generate_charts(project, projections)
        for buf in charts:
            cls._add_picture_safely(doc, buf, width=Inches(5.5))
            doc.add_paragraph() # Spacer

    @classmethod
    def _add_section_R_liquidity(cls, doc, project, projections, theme=None):
        """R. Liquidity Analysis."""
        doc.add_heading("SECTION R: LIQUIDITY & WORKING CAPITAL", level=1)
        if not projections: return
        table = doc.add_table(rows=1, cols=len(projections)+1)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        headers = ["Liquidity Ratios / FY"] + cls._get_year_headers(projections, project)
        shd = theme.PRIMARY_HEX[1:] if theme else "0D47A1"
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cls._set_cell_background(cell, shd)
 
        
        rows = [("Current Assets", "current_assets"), ("Current Ratio", "current_ratio")]
        for label, key in rows:
            row = table.add_row().cells
            row[0].text = label
            for i, p in enumerate(projections):
                row[i+1].text = cls._fmt_val(p.get(key, 0))

    @classmethod
    def _add_section_U_margin(cls, doc, project, projections):
        """U. Security Margin Analysis."""
        h = doc.add_heading("SECTION U: SECURITY MARGIN ANALYSIS", level=1)
        cls._set_keep_with_next(h)
        total_assets = sum(a.cost for a in project.assets)
        total_loan = project.loan.term_loan_amount
        margin = ( (total_assets - total_loan) / total_assets * 100 ) if total_assets > 0 else 0
        
        doc.add_paragraph(f"Total Fixed Assets: Rs. {cls._fmt_val(total_assets)} Lakhs")
        doc.add_paragraph(f"Proposed Bank Finance: Rs. {cls._fmt_val(total_loan)} Lakhs")
        doc.add_paragraph(f"Effective Promoter Margin: {cls._fmt_val(margin)}%")
        doc.add_paragraph("The project maintains healthy margins as per standard banking conservative guidelines.")

    @classmethod
    def _add_section_V_repayment(cls, doc, project):
        """V. Repayment Schedule Summary."""
        h = doc.add_heading("SECTION V: LOAN REPAYMENT SUMMARY", level=1)
        cls._set_keep_with_next(h)
        l = project.loan
        doc.add_paragraph(f"Requested Facility: {l.facility_type}")
        doc.add_paragraph(f"Loan Amount: Rs. {cls._fmt_val(l.term_loan_amount)} Lakhs")
        doc.add_paragraph(f"Proposed Tenure: {l.term_loan_tenure_years} Years ({l.term_loan_tenure_years*12} Months)")
        doc.add_paragraph(f"Moratorium Period: {project.assumptions.moratorium_months} Months")
        doc.add_paragraph(f"Interest Rate: {project.assumptions.interest_on_tl}% p.a. (Indicative)")

    @classmethod
    def _add_section_AA_monthly_repayment(cls, doc, project, theme=None):
        """AA. Monthly Repayment List."""
        h = doc.add_heading("SECTION AA: PROJECTED MONTHLY REPAYMENT", level=1)
        cls._set_keep_with_next(h)
        doc.add_paragraph("Indicative repayment schedule based on requested tenure and interest rates.")
        
        # Compact table for monthly data
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        headers = ["Month", "Opening", "EMI / Installment", "Closing"]
        shd = theme.PRIMARY_HEX[1:] if theme else "0D47A1"
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cls._set_cell_background(cell, shd)

        amt = project.loan.term_loan_amount
        months = project.loan.term_loan_tenure_years * 12
        installment = amt / months if months > 0 else 0
        
        for m in range(1, min(13, months + 1)): # Show first year
            row = table.add_row().cells
            row[0].text = str(m)
            row[1].text = cls._fmt_val(amt - (m-1)*installment)
            row[2].text = cls._fmt_val(installment)
            row[3].text = cls._fmt_val(amt - m*installment)
        if months > 12:
            doc.add_paragraph("... schedule continues for remaining tenure.")

    @classmethod
    def _add_section_W_cma_data(cls, doc, project, projections, theme=None):
        """W. CMA Data key mathematical extracts."""
        h = doc.add_heading("SECTION W: CMA DATA PRESENTATION", level=1)
        cls._set_keep_with_next(h)
        doc.add_paragraph("Consolidated financial mapping for Credit Monitoring Arrangement (CMA) appraisal.")
        
        if not projections: return
        headers = ["CMA KEY EXTRACTS"] + cls._get_year_headers(projections, project)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        
        # Style Header
        shd = theme.PRIMARY_HEX[1:] if theme else "0D47A1"
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cls._set_cell_background(cell, shd)

        rows = [
            ("Adjusted Tangible Net Worth", lambda r: r.get("share_capital",0) + r.get("reserves_surplus",0)),
            ("Working Capital Gap", lambda r: r.get("current_assets",0) - r.get("creditors",0)),
            ("MPBF (Method II)", lambda r: (r.get("current_assets",0)*0.75 - r.get("creditors",0))),
            ("Current Ratio (Assessed)", lambda r: r.get("current_ratio", 0)),
        ]
        for label, formula in rows:
            r_cells = table.add_row().cells
            r_cells[0].text = label
            r_cells[0].paragraphs[0].runs[0].bold = True
            for i, p in enumerate(projections):
                r_cells[i+1].text = cls._fmt_val(formula(p))

    @classmethod
    def _add_section_X_assumptions(cls, doc, project):
        """X. Financial Assumptions & Institutional Notes."""
        doc.add_heading("SECTION X: NOTES & FINANCIAL ASSUMPTIONS", level=1)
        doc.add_paragraph("The projections are based on the following key managerial and economic assumptions:")
        
        ass = project.assumptions
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        hdr = table.rows[0].cells
        hdr[0].text = "Particulars / Parameter"
        hdr[1].text = "Assumed Basis"
        for c in hdr: c.paragraphs[0].runs[0].bold = True
        
        data = [
            ("Sales Growth Rate", f"{ass.sales_growth_percent}% p.a. (Compounded)"),
            ("Gross Profit Margin", f"{ass.gp_percent}% on Sales"),
            ("Tax Provision Rate", f"{ass.tax_rate_percent}% on PBT"),
            ("Depreciation Policy", f"{ass.depreciation_method} Basis"),
            ("Interest on Term Loan", f"{ass.interest_on_tl}% p.a."),
            ("Interest on WC Finance", f"{ass.interest_on_cc}% p.a."),
        ]
        for label, val in data:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = val

        doc.add_paragraph() # Spacer
        
        # Institutional Notes Heading (Styled like PDF)
        p = doc.add_paragraph("Notes to the Project Report")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.bold = True
        run.font.size = Pt(12)
        cls._set_keep_with_next(p)
        
        notes = [
            ("a.", "Depreciation has been computed in accordance with the depreciation rates prescribed in the Income Tax Act. A separate depreciation schedule has been provided for reference and calculation purposes."),
            ("b.", "The data presented, including sensitivity analysis and balance sheet synopsis, has been prepared utilizing standard financial assumptions and calculations."),
            ("c.", "The financial projections and assessments are based on the assumption that there will be no changes in government policies and rules that may impact the loan applicant's business. Furthermore, it is assumed that no abnormal events will occur during the lifespan of the project or business."),
            ("d.", "Provision for Income Tax has been made on the Rules and Regulations which are applicable for current scenario."),
            ("e.", "The financial statements have been prepared under the standard assumption that the fiscal year-end occurs in March."),
            ("f.", "The details of indirect expenses, break-even analysis, and security margin calculation have been provided in separate annexures for reference."),
            ("g.", "The financial data pertaining to revenue from business operations, asset additions, existing obligations, etc., has been presented based on the information provided by the client."),
            ("h.", "The projected data included in this report represents future-oriented financial information. It has been prepared based on the best judgment of the applicants, incorporating assumptions regarding the most probable set of economic conditions. However, it is important to note that this information should not be considered as a forecast."),
            ("i.", "The information pertaining to the business entity, owner's profile, employment details, feasibility studies, industry analysis, market potential, current scenario, and challenges/solutions has been compiled based on discussions and inputs provided by the loan applicant.")
        ]
        
        for char, text in notes:
            # Bullet-like table for clean alignment
            n_tbl = doc.add_table(rows=1, cols=2)
            # Remove table borders for the notes list
            for row in n_tbl.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for b in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f'w:{b}')
                        border.set(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}val', 'nil')
                        tcBorders.append(border)
                    tcPr.append(tcBorders)
            
            cls._prevent_table_split(n_tbl)
            cells = n_tbl.rows[0].cells
            # Set relative widths
            cells[0].width = Inches(0.4)
            cells[1].width = Inches(6.0)
            
            cells[0].text = char
            cells[1].text = text
            cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    @classmethod
    def _add_section_A_cover(cls, doc, project, theme=None):
        """A. Professional Word Cover Page — Mode-specific."""
        from services.cma.models import ReportMode, SchemeType
        report_mode = project.profile.report_mode
        scheme = project.profile.scheme_type
        
        primary_rgb = RGBColor(*theme.PRIMARY_RGB) if theme else RGBColor(13, 71, 161)
        secondary_rgb = RGBColor(*theme.SECONDARY_RGB) if theme else RGBColor(100, 100, 100)
        
        # 1. Main Title Logic
        if report_mode == ReportMode.CMA.value:
            doc_title = "BANKER'S CMA DATA & ANALYSIS"
            sub_title = "WORKING CAPITAL ASSESSMENT REPORT"
        elif report_mode == ReportMode.LITE.value:
            doc_title = "PROJECT REPORT (SUMMARY)"
            sub_title = f"{project.profile.loan_type} PROPOSAL"
        else:
            doc_title = "DETAILED PROJECT REPORT (DPR)"
            sub_title = "PREMIUM BANK SUBMISSION PACK"

        # Override based on Scheme
        if scheme == SchemeType.MUDRA.value:
            doc_title = "PRADHAN MANTRI MUDRA YOJANA (PMMY)"
            sub_title = "PROJECT REPORT FOR MUDRA LOAN"
        elif scheme == SchemeType.PMEGP.value:
            doc_title = "PMEGP PROJECT REPORT"
            sub_title = "PRIME MINISTER'S EMPLOYMENT GENERATION PROGRAMME"
        
        t = doc.add_paragraph(doc_title)
        t.paragraph_format.space_before = Pt(60)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.runs[0] if t.runs else t.add_run()
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = primary_rgb
        
        st = doc.add_paragraph(sub_title)
        st.paragraph_format.space_after = Pt(40)
        st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_st = st.runs[0] if st.runs else st.add_run()
        run_st.font.size = Pt(14)
        run_st.font.color.rgb = secondary_rgb
        
        ent = doc.add_paragraph(project.profile.business_name.upper() or "BUSINESS NAME")
        ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ent.runs[0] if ent.runs else ent.add_run()
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = primary_rgb
        
        addr = doc.add_paragraph(project.profile.address or "Business Address Line")
        addr.paragraph_format.space_after = Pt(80)
        addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph(f"Report Mode: {report_mode} | {datetime.now().strftime('%B %d, %Y')}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].italic = True
               # 2. Branding Section
        b = project.branding
        if b.prepared_by or b.firm_name:
            doc.add_paragraph()
            p = doc.add_paragraph("REPORT PREPARED BY")
            p.paragraph_format.space_before = Pt(80)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(10)
            run.font.color.rgb = secondary_rgb
            
            if b.prepared_by:
                p = doc.add_paragraph(b.prepared_by.upper())
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0] if p.runs else p.add_run()
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = primary_rgb

            p = doc.add_paragraph( (b.firm_name or "Professional Services").upper() )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run()
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = primary_rgb
            
            if b.contact_line:
                p = doc.add_paragraph(b.contact_line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0] if p.runs else p.add_run()
                run.font.size = Pt(11)

    @classmethod
    def _add_section_B_contents(cls, doc, project):
        """B. Table of Contents (Manual for Word Parity)."""
        from services.cma.models import ReportMode
        doc.add_heading("Table of Contents", level=1)
        mode = project.profile.report_mode
        
        sections = [
            ("A", "PROJECT COVER PAGE"),
            ("B", "TABLE OF CONTENTS"),
            ("C", "EXECUTIVE SUMMARY"),
            ("DASH", "EXECUTIVE DASHBOARD"),
            ("D", "PROJECT SNAPSHOT"),
            ("E", "ENTITY PROFILE"),
            ("F", "PROMOTER DETAILS"),
            ("G", "EMPLOYMENT DETAILS"),
            ("H", "COST OF PROJECT"),
            ("I", "MEANS OF FINANCE"),
            ("J", "FINANCIAL OVERVIEW"),
            ("K", "GRAPHICAL ANALYSIS"),
            ("L", "OPERATING STATEMENT"),
            ("M", "BALANCE SHEET"),
            ("N", "CASH FLOW STATEMENT"),
            ("O", "FIXED ASSETS SCHEDULE"),
            ("V", "LOAN REPAYMENT SCHEDULE"),
            ("AA", "MONTHLY REPAYMENT"),
            ("Q", "DSCR ANALYSIS"),
            ("P", "EXPENSES BREAKDOWN"),
            ("R", "LIQUIDITY ANALYSIS"),
            ("S", "SENSITIVITY ANALYSIS"),
            ("T", "BREAK-EVEN ANALYSIS"),
            ("W", "CMA DATA EXTRACTS"),
            ("AC", "MPBF ASSESSMENT"),
            ("AD", "BANK-READINESS"),
            ("X", "NOTES & ASSUMPTIONS"),
            ("Y", "SECURITY DETAILS"),
            ("Z", "FINAL DECLARATION"),
        ]
        
        # Filter based on mode (simplified for parity)
        for code, title in sections:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.add_run(f"Section {code}: ").bold = True
            p.add_run(title)

    @classmethod
    def _clean_text(cls, text: str) -> str:
        """Removes raw HTML tags like <b> from the narrative text for Word."""
        import re
        if not text: return ""
        return re.sub(r'<[^>]*>', '', text)

    @classmethod
    def _add_section_C_summary(cls, doc, project):
        """C. Executive Summary."""
        doc.add_heading("SECTION C: EXECUTIVE SUMMARY", level=1)
        text = cls._clean_text(NarrativeService.generate_section("executive_summary", project))
        doc.add_paragraph(text)
        
        doc.add_heading("PROJECT RATIONALE", level=2)
        rationale = cls._clean_text(NarrativeService.generate_section("project_rationale", project))
        doc.add_paragraph(rationale)

    @classmethod
    def _add_section_D_snapshot(cls, doc, project, theme=None):
        """D. Project Snapshot Table."""
        h = doc.add_heading("SECTION D: PROJECT SNAPSHOT", level=1)
        cls._set_keep_with_next(h)
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        total_assets = sum(a.cost for a in project.assets)
        total_cost = total_assets + project.loan.working_capital_requirement
        
        data = [
            ("Proposed Business Name", project.profile.business_name),
            ["Project Opportunity", project.loan.purpose or "Capacity Expansion / Scaling"],
            ["Business Model", project.profile.business_mode],
            ("Legal Entity Constitution", project.profile.entity_type),
            ("Primary Industry Category", project.profile.business_category),
            ("Administrative Location", project.profile.address),
            ("Commencement / Est. Date", project.profile.establishment_date),
            ("Total Appraised Project Cost", f"Rs. {cls._fmt_val(total_cost)} Lakhs"),
            ("Total Credit Assistance Sought", f"Rs. {cls._fmt_val(project.loan.term_loan_amount + project.loan.cash_credit_amount)} Lakhs"),
        ]
        for label, val in data:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = val
            row[0].paragraphs[0].runs[0].bold = True
            cls._set_cell_background(row[0], "F2F2F2")

    @classmethod
    def _add_section_E_entity(cls, doc, project, theme=None):
        """E. Entity Profile."""
        h = doc.add_heading("SECTION E: ENTITY PROFILE", level=1)
        cls._set_keep_with_next(h)
        text = cls._clean_text(NarrativeService.generate_section("business_overview", project))
        doc.add_paragraph(text)
        
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        data = [
            ("PAN / Income Tax ID", project.profile.pan),
            ("Entity Structure", project.profile.entity_type),
            ("Full Operational Address", project.profile.address),
        ]
        for label, val in data:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = val
            row[0].paragraphs[0].runs[0].bold = True

    @classmethod
    def _add_section_F_promoter(cls, doc, project):
        """F. Promoter Profile."""
        doc.add_heading("SECTION F: PROMOTER PROFILE & MANAGEMENT", level=1)
        text = cls._clean_text(NarrativeService.generate_section("promoter_profile", project))
        doc.add_paragraph(text)

    @classmethod
    def _add_section_G_employment(cls, doc, project, theme=None):
        """G. Employment & Manpower."""
        h = doc.add_heading("SECTION G: EMPLOYMENT & MANPOWER", level=1)
        cls._set_keep_with_next(h)
        text = cls._clean_text(NarrativeService.generate_section("employment_details", project))
        doc.add_paragraph(text)
        
        count = project.profile.employee_count
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        hdr = table.rows[0].cells
        hdr[0].text = "Category"
        hdr[1].text = "Headcount"
        hdr[0].paragraphs[0].runs[0].bold = True
        hdr[1].paragraphs[0].runs[0].bold = True
        
        row = table.add_row().cells
        row[0].text = "Total Workforce Requirements"
        row[1].text = str(count)

    @classmethod
    def _add_section_H_cost(cls, doc, project, theme=None):
        """H. Cost of Project Table."""
        h = doc.add_heading("SECTION H: DETAILED COST OF PROJECT", level=1)
        cls._set_keep_with_next(h)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        shd = theme.PRIMARY_HEX[1:] if theme else "0D47A1"
        hdr = table.rows[0].cells
        col_names = ["Sr.", "Particulars / Asset Item", "Amount (Rs. Lakhs)"]
        for i, text in enumerate(col_names):
            cell = hdr[i]
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cls._set_cell_background(cell, shd)
        
        total = 0
        for i, asset in enumerate(project.assets):
            row = table.add_row().cells
            row[0].text = str(i + 1)
            row[1].text = asset.name
            row[2].text = cls._fmt_val(asset.cost)
            total += asset.cost
        
        row = table.add_row().cells
        row[1].text = "Working Capital Requirement"
        row[2].text = cls._fmt_val(project.loan.working_capital_requirement)
        total += project.loan.working_capital_requirement
        
        row = table.add_row().cells
        row[1].text = "TOTAL ESTIMATED PROJECT COST"
        row[2].text = cls._fmt_val(total)
        row[1].paragraphs[0].runs[0].bold = True
        row[2].paragraphs[0].runs[0].bold = True

    @classmethod
    def _add_section_I_finance(cls, doc, project, theme=None):
        """I. Means of Finance Table."""
        h = doc.add_heading("SECTION I: PROPOSED MEANS OF FINANCE", level=1)
        cls._set_keep_with_next(h)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        shd = theme.PRIMARY_HEX[1:] if theme else "0D47A1"
        hdr = table.rows[0].cells
        col_names = ["Sr.", "Source of Finance", "Amount (Rs. Lakhs)"]
        for i, text in enumerate(col_names):
            cell = hdr[i]
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cls._set_cell_background(cell, shd)
        
        total_cost = sum(a.cost for a in project.assets) + project.loan.working_capital_requirement
        tl = project.loan.term_loan_amount
        cc = project.loan.cash_credit_amount
        promoter = total_cost - (tl + cc)
        
        sources = [
            ("Promoter's Capital / Equity", promoter),
            ("Proposed Term Loan", tl),
            ("Proposed Working Capital (CC/OD)", cc),
        ]
        for i, (name, amt) in enumerate(sources):
            row = table.add_row().cells
            row[0].text = str(i+1)
            row[1].text = name
            row[2].text = cls._fmt_val(amt)
            
        row = table.add_row().cells
        row[1].text = "TOTAL MEANS OF FINANCE"
        row[2].text = cls._fmt_val(total_cost)
        row[1].paragraphs[0].runs[0].bold = True
        row[2].paragraphs[0].runs[0].bold = True

    @classmethod
    def _add_section_J_financial_data(cls, doc, project, projections, theme=None):
        """J. Financial Overview."""
        h = doc.add_heading("SECTION J: FINANCIAL OVERVIEW & CORE RATIOS", level=1)
        cls._set_keep_with_next(h)
        table = doc.add_table(rows=1, cols=len(projections)+1)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        headers = ["Indicators / FY"] + cls._get_year_headers(projections, project)
        shd = theme.PRIMARY_HEX[1:] if theme else "0D47A1"
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cls._set_cell_background(cell, shd)
        
        metrics = [("Annual Revenue", "revenue"), ("Profit After Tax (PAT)", "pat"), ("DSCR", "dscr"), ("Current Ratio", "current_ratio")]
        for label, key in metrics:
            row = table.add_row().cells
            row[0].text = label
            row[0].paragraphs[0].runs[0].bold = True
            for i, p in enumerate(projections):
                row[i+1].text = cls._fmt_val(p.get(key, 0))

    @classmethod
    def _add_section_L_operating_stmt(cls, doc, project, projections, theme=None):
        """L. Projected Operating Statement."""
        h = doc.add_heading("SECTION L: PROJECTED OPERATING STATEMENT", level=1)
        cls._set_keep_with_next(h)
        if not projections: return
        table = doc.add_table(rows=1, cols=len(projections)+1)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        headers = ["Particulars (Rs. Lakhs)"] + cls._get_year_headers(projections, project)
        shd = theme.PRIMARY_HEX[1:] if theme else "0D47A1"
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cls._set_cell_background(cell, shd)
        
        metrics = [
            ("Gross Revenue / Sales", "revenue"),
            ("GROSS PROFIT (GP)", "gp_amt"),
            ("Indirect Expenses", "ind_exp"),
            ("EBITDA", "ebitda"),
            ("Depreciation", "depreciation"),
            ("Interest on Loans", "tl_interest"),
            ("PROFIT AFTER TAX (PAT)", "pat"),
        ]
        for label, key in metrics:
            row = table.add_row().cells
            row[0].text = label
            for i, p in enumerate(projections):
                row[i+1].text = cls._fmt_val(p.get(key, 0))
            if "PAT" in label or "GROSS PROFIT" in label:
                row[0].paragraphs[0].runs[0].bold = True

    @classmethod
    def _add_section_M_balance_sheet(cls, doc, project, projections, theme=None):
        """M. Projected Balance Sheet."""
        h = doc.add_heading("SECTION M: PROJECTED BALANCE SHEET", level=1)
        cls._set_keep_with_next(h)
        if not projections: return
        table = doc.add_table(rows=1, cols=len(projections)+1)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        headers = ["Liabilities & Assets"] + cls._get_year_headers(projections, project)
        shd = theme.PRIMARY_HEX[1:] if theme else "0D47A1"
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cls._set_cell_background(cell, shd)
        
        rows = [
            ("Share Capital / Net Worth", "reserves_surplus"),
            ("Long Term Bank Borrowings", "term_loan_bal"),
            ("Current Liabilities", "creditors"),
            ("TOTAL LIABILITIES", "total_liabilities"),
            ("-", ""),
            ("Fixed Assets (Net Block)", "net_block"),
            ("Current Assets", "current_assets"),
            ("TOTAL ASSETS", "total_assets"),
        ]
        for label, key in rows:
            if label == "-":
                table.add_row()
                continue
            row = table.add_row().cells
            row[0].text = label
            for i, p in enumerate(projections):
                row[i+1].text = cls._fmt_val(p.get(key, 0))
            if "TOTAL" in label:
                row[0].paragraphs[0].runs[0].bold = True

    @classmethod
    def _add_section_N_cash_flow(cls, doc, project, projections, theme=None):
        """N. Projected Cash Flow Statement."""
        h = doc.add_heading("SECTION N: PROJECTED CASH FLOW STATEMENT", level=1)
        cls._set_keep_with_next(h)
        if not projections: return
        table = doc.add_table(rows=1, cols=len(projections)+1)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        hdr = table.rows[0].cells
        hdr[0].text = "Source / Use of Funds"
        labels = cls._get_year_headers(projections, project)
        for i, lbl in enumerate(labels): hdr[i+1].text = lbl
        for c in hdr: c.paragraphs[0].runs[0].bold = True
        
        rows = [
            ("Net Profit After Tax", "pat"),
            ("Add: Depreciation", "depreciation"),
            ("A. Fund from Operations", "cash_accruals"),
            ("Less: Loan Repayment", "tl_repayment"),
            ("Net Cash Surplus / (Deficit)", "net_cf"),
        ]
        for label, key in rows:
            row = table.add_row().cells
            row[0].text = label
            for i, p in enumerate(projections):
                # Use .get() for safety and recalculate net_cf if it's based on tl_repayment
                ca = p.get('cash_accruals', 0)
                tlr = p.get('tl_repayment', 0)
                
                if key == "net_cf": 
                    val = ca - tlr
                else: 
                    val = p.get(key, 0)
                
                row[i+1].text = cls._fmt_val(val)

    @classmethod
    def _add_section_O_fixed_assets(cls, doc, project, projections, theme=None):
        """O. Fixed Assets & Depreciation Annexure — Year-Wise Movement."""
        h = doc.add_heading("SECTION O: FIXED ASSETS & DEPRECIATION SCHEDULE", level=1)
        cls._set_keep_with_next(h)
        doc.add_paragraph(f"Method: {project.assumptions.depreciation_method} | (Rs. in Lakhs)")
        
        if not projections: return
        
        # Table Structure: Year | Opening | Additions | Depreciation | Closing
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        hdr = table.rows[0].cells
        hdr[0].text = "Fiscal Year"
        hdr[1].text = "Opening Bal"
        hdr[2].text = "Additions"
        hdr[3].text = "Depreciation"
        hdr[4].text = "Net Block"
        for c in hdr: c.paragraphs[0].runs[0].bold = True
        
        for p in projections:
            lbl = p.get("year_label", "N/A")
            is_act = p.get("is_actual", False)
            
            opening = p.get("opening_fixed_assets", 0.0)
            additions = p.get("fixed_asset_additions", 0.0)
            depreciation = p.get("depreciation", 0.0)
            closing = p.get("net_fixed_assets", 0.0)
            
            # Historical fallback
            if is_act and opening == 0 and additions == 0:
                opening = closing + depreciation
                
            row = table.add_row().cells
            row[0].text = f"{lbl} {'(ACT)' if is_act else '(PROJ)'}"
            row[1].text = f"{cls._fmt_val(opening)}"
            row[2].text = f"{cls._fmt_val(additions)}"
            row[3].text = f"{cls._fmt_val(depreciation)}"
            row[4].text = f"{cls._fmt_val(closing)}"
            
        doc.add_paragraph("\nNote: Depreciation is calculated as per standard income tax rates and WDV/SLM method as applicable to the industry.")

    @classmethod
    def _add_section_P_expenses(cls, doc, project, projections, theme=None):
        """P. Indirect Expenses Breakdown."""
        h = doc.add_heading("SECTION P: INDIRECT EXPENSES BREAKDOWN", level=1)
        cls._set_keep_with_next(h)
        if not projections: return
        table = doc.add_table(rows=1, cols=len(projections)+1)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        headers = ["Cash Movement Component"] + cls._get_year_headers(projections, project)
        shd = theme.PRIMARY_HEX[1:] if theme else "0D47A1"
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cls._set_cell_background(cell, shd)
        
        heads = ["Salary & Wages", "Power & Fuel", "Rent & Rates", "Admin & Misc"]
        for head in heads:
            row = table.add_row().cells
            row[0].text = head
            for i, p in enumerate(projections):
                # FIX: Remove 'detailed' key, access top-level dict
                eb = p.get("expense_breakdown", {})
                val = eb.get(head, 0.0)
                row[i+1].text = cls._fmt_val(val)

    @classmethod
    def _add_section_Q_dscr(cls, doc, project, projections, theme=None):
        """Q. DSCR Analysis Table."""
        h = doc.add_heading("SECTION Q: DEBT SERVICE COVERAGE RATIO (DSCR)", level=1)
        cls._set_keep_with_next(h)
        if not projections: return
        table = doc.add_table(rows=1, cols=len(projections)+1)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        hdr = table.rows[0].cells
        hdr[0].text = "Particulars"
        labels = cls._get_year_headers(projections, project)
        for i, lbl in enumerate(labels): hdr[i+1].text = lbl
        for c in hdr: c.paragraphs[0].runs[0].bold = True
        
        rows = [("Fund Available", "cash_accruals"), ("Repayment Obligation", "tl_repayment"), ("DSCR", "dscr")]
        for label, key in rows:
            row = table.add_row().cells
            row[0].text = label
            for i, p in enumerate(projections):
                row[i+1].text = cls._fmt_val(p.get(key, 0))
            if "DSCR" in label:
                row[0].paragraphs[0].runs[0].bold = True

    @classmethod
    def _add_section_S_sensitivity(cls, doc, project, projections, theme=None):
        """S. Sensitivity Analysis."""
        h = doc.add_heading("SECTION S: SENSITIVITY TEST ANALYSIS", level=1)
        cls._set_keep_with_next(h)
        if not projections: return
        table = doc.add_table(rows=1, cols=len(projections)+1)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        headers = ["Sensitivity Component"] + cls._get_year_headers(projections, project)
        shd = theme.PRIMARY_HEX[1:] if theme else "0D47A1"
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cls._set_cell_background(cell, shd)
        
        scenarios = [("Base Case DSCR", "dscr"), ("-10% Revenue Shortfall", "minus_10pct"), ("-20% Revenue Shortfall", "minus_20pct")]
        for label, key in scenarios:
            row = table.add_row().cells
            row[0].text = label
            for i, p in enumerate(projections):
                if key == "dscr": val = p.get('dscr', 0)
                else: val = p.get("sensitivity", {}).get(key, 0)
                row[i+1].text = cls._fmt_val(val)

    @classmethod
    def _add_section_T_bep(cls, doc, project, projections, theme=None):
        """T. Break-Even Analysis."""
        h = doc.add_heading("SECTION T: BREAK-EVEN POINT (BEP) TARGETS", level=1)
        cls._set_keep_with_next(h)
        if not projections: return
        table = doc.add_table(rows=1, cols=len(projections)+1)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        hdr = table.rows[0].cells
        hdr[0].text = "Particulars"
        labels = cls._get_year_headers(projections, project)
        for i, lbl in enumerate(labels): hdr[i+1].text = lbl
        for c in hdr: c.paragraphs[0].runs[0].bold = True
        
        rows = [("Total Fixed Costs", "fixed_costs"), ("Break-Even Sales", "bep_sales")]
        for label, key in rows:
            row = table.add_row().cells
            row[0].text = label
            for i, p in enumerate(projections):
                row[i+1].text = cls._fmt_val(p.get(key, 0))

    @classmethod
    def _add_section_Y_security(cls, doc, project):
        """Y. Security Details."""
        doc.add_heading("SECTION Y: SECURITY & COLLATERAL", level=1)
        p1 = doc.add_paragraph()
        run1 = p1.add_run("Primary Security: ")
        run1.bold = True
        p1.add_run("Personal guarantee of promoters and hypothecation of all assets created out of bank finance including Plant, Machinery, Stock and Receivables.")
        
        p2 = doc.add_paragraph()
        run2 = p2.add_run("Collateral Security: ")
        run2.bold = True
        p2.add_run("Additional tangible security as per the bank's specific sanction terms.")

    @classmethod
    def _add_section_Z_declaration(cls, doc, project):
        """Z. Final Declaration."""
        doc.add_heading("SECTION Z: DECLARATION & CERTIFICATION", level=1)
        doc.add_paragraph("I/We hereby declare that all information provided in this Integrated Financial Study / Detailed Project Report is true and accurate to the best of my/our knowledge.")
        
        for _ in range(4): doc.add_paragraph()
        table = doc.add_table(rows=1, cols=2)
        cls._prevent_table_split(table)
        cells = table.rows[0].cells
        cells[0].text = f"Prepared by:\nDate: {datetime.now().strftime('%d/%m/%Y')}"
        cells[1].text = f"For {project.profile.business_name.upper()}\n\nAuthorized Signatory"
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    @classmethod
    def _add_section_AC_mpbf(cls, doc, project, projections, theme=None):
        """AC. MPBF Analysis."""
        doc.add_heading("SECTION AC: MPBF & WORKING CAPITAL ASSESSMENT", level=1)
        from services.cma.mpbf_service import MpbfService
        res = MpbfService.calculate_mpbf(project, projections)
        if res["status"] in ["error", "insufficient"]:
            doc.add_paragraph(f"⚠️ {res.get('message', 'Data insufficient for MPBF calculation.')}")
            return
        
        doc.add_paragraph(f"Turnover Assessed: Rs. {cls._fmt_val(res.get('turnover', 0))} Lakhs")
        doc.add_paragraph(f"Requested Limit: Rs. {cls._fmt_val(res.get('requested_limit', 0))} Lakhs")
        doc.add_paragraph(f"Permissible Limit (MPBF): Rs. {cls._fmt_val(res.get('permissible_limit', 0))} Lakhs")
        doc.add_paragraph(f"Status: {res.get('risk_level', '')}")
        if res.get("suggested_correction") and res["suggested_correction"] != "None needed.":
            doc.add_paragraph(f"Advice: {res['suggested_correction']}")
        if res.get("shortfall_nwc", 0) > 0:
            doc.add_paragraph(f"NWC Shortfall: Rs. {cls._fmt_val(res.get('shortfall_nwc', 0))} Lakhs")

    @classmethod
    def _add_section_AD_readiness(cls, doc, project, projections, theme=None):
        """AD. Bank-Readiness Check."""
        from services.cma.readiness_service import ReadinessService
        res = ReadinessService.evaluate_readiness(project, projections)
        
        if theme and theme.mode_key == "lite":
            doc.add_heading("READINESS SNAPSHOT", level=1)
            doc.add_paragraph(f"Overall Readiness Level: {res['readiness_level']}")
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            cls._prevent_table_split(table)
            hdr = table.rows[0].cells
            hdr[0].text = "Parameter"
            hdr[1].text = "Current Value"
            for c in hdr: c.paragraphs[0].runs[0].bold = True
            for chk in res["checks"]:
                row = table.add_row().cells
                row[0].text = chk["name"]
                row[1].text = str(chk["value"])
        else:
            h = doc.add_heading("SECTION AD: BANK-READINESS HEALTH CHECK", level=1)
            cls._set_keep_with_next(h)
            doc.add_paragraph(f"Overall Readiness Level: {res['readiness_level']}")
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            cls._prevent_table_split(table)
            hdr = table.rows[0].cells
            hdr[0].text = "Parameter"
            hdr[1].text = "Value"
            hdr[2].text = "Assessment & Advice"
            for c in hdr: c.paragraphs[0].runs[0].bold = True
            
            for c in res["checks"]:
                row = table.add_row().cells
                row[0].text = c["name"]
                row[1].text = c["value"]
                row[2].text = c["advice"]
    @classmethod
    def _add_section_DASH_dashboard(cls, doc, project, projections, theme=None):
        """DASH. Feasibility Dashboard Table."""
        doc.add_heading("PROJECT FEASIBILITY DASHBOARD", level=1)
        if not projections: return
        
        from services.cma.readiness_service import ReadinessService
        audit = ReadinessService.evaluate_readiness(project, projections)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        cls._prevent_table_split(table)
        hdr = table.rows[0].cells
        hdr[0].text = "Key Performance Indicator"
        hdr[1].text = "Value"
        hdr[2].text = "Benchmark"
        hdr[3].text = "Status"
        for c in hdr: c.paragraphs[0].runs[0].bold = True
        
        proj_only = [r for r in projections if not r.get('is_actual')]
        n_proj = max(1, len(proj_only))
        avg_dscr = sum(r.get('dscr', 0) for r in proj_only) / n_proj
        peak_rev = max(r.get('revenue', 0) for r in projections)
        avg_pat_pct = sum( (r.get('pat',0)/r.get('revenue',1)*100) for r in projections if r.get('revenue',0)>0 ) / max(1, len(projections))
        last_cr = projections[-1].get('current_ratio', 0)

        # Sync Status Labels
        def get_status(metric_name, fallback_pass):
            for c in audit["checks"]:
                if metric_name in c["name"]:
                    if c["level"] == "PASS": return str(fallback_pass)
                    if c["level"] == "CRITICAL": 
                        return "MARGIN WATCH" if "Margin" in metric_name else "REVISE"
                    return "BORDERLINE"
            return str(fallback_pass)

        dash_data = [
            ("DEBT REPAYMENT CAPACITY (DSCR)", f"{cls._fmt_val(avg_dscr)}", "> 1.25", get_status("DSCR", "HEALTHY")),
            ("PEAK PROJECTED REVENUE", f"Rs. {cls._fmt_val(peak_rev)} L", "Targeted Capacity", "OPTIMIZED"),
            ("AVG. NET PROFIT MARGIN (%)", f"{cls._fmt_val(avg_pat_pct)}%", "10-15%", get_status("Profit Margin", "HEALTHY")),
            ("CURRENT LIQUIDITY RATIO", f"{cls._fmt_val(last_cr)}", "> 1.17", get_status("Current Ratio", "STABLE"))
        ]
        
        for kpi, val, bench, status in dash_data:
            row = table.add_row().cells
            row[0].text = str(kpi)
            row[1].text = str(val)
            row[2].text = str(bench)
            row[3].text = str(status)
            row[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()
        h = doc.add_heading("STRATEGIC PROJECT HIGHLIGHTS:", level=2)
        cls._set_keep_with_next(h)
        
        cash_flow_desc = "TIGHT" if avg_dscr < 1.20 else "ADEQUATE" if avg_dscr < 1.40 else "STRONG"
        
        highlights = [
            f"Capacity Utilization: The project targets a peak revenue of Rs. {cls._fmt_val(peak_rev)} Lakhs reflecting 80-90% efficiency.",
            f"Debt Comfort: With an average DSCR of {cls._fmt_val(avg_dscr)}, the enterprise shows {cash_flow_desc} cash flow for bank obligations.",
            f"Operational Viability: The Net Profit margins are projected to stabilize at {cls._fmt_val(avg_pat_pct)}% following the startup phase."
        ]
        for i, h_text in enumerate(highlights):
            pk = doc.add_paragraph(h_text, style='List Bullet')
            # Keep bullets together except last one
            if i < len(highlights) - 1:
                cls._set_keep_with_next(pk)

        doc.add_paragraph()
        p = doc.add_paragraph(f"Overall Assessment: {audit['readiness_level']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(10)

